"""
生成清数智算股票分析算法文档 (Word .docx)
"""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import os

doc = Document()

# ── 样式设置 ──
style = doc.styles['Normal']
font = style.font
font.name = '微软雅黑'
font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.5

# 标题样式
for i in range(1, 4):
    heading_style = doc.styles[f'Heading {i}']
    heading_style.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)

# ── 封面 ──
doc.add_paragraph()
doc.add_paragraph()
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('清数智算投研工作台\n股票分析算法技术文档')
run.font.size = Pt(28)
run.font.bold = True
run.font.color.rgb = RGBColor(0x0F, 0x3D, 0x6E)

doc.add_paragraph()
subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('清数智算创新团队\n版本 2.0 · 2026年6月')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x47, 0x55, 0x69)

doc.add_paragraph()
disclaimer = doc.add_paragraph()
disclaimer.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = disclaimer.add_run('免责声明：本文档描述的所有算法仅用于研究和辅助分析，不构成任何投资建议。\n股市有风险，投资需谨慎。')
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x94, 0x3A, 0x3F)

doc.add_page_break()

# ── 目录 ──
doc.add_heading('目录', level=1)
toc_items = [
    '1. 系统架构概述',
    '2. 数据采集与预处理',
    '3. 技术指标计算',
    '4. 五维评分模型',
    '  4.1 质量评分（30分）',
    '  4.2 估值评分（20分）',
    '  4.3 成长评分（20分）',
    '  4.4 趋势评分（20分）',
    '  4.5 风险评分（10分）',
    '  4.6 综合评级',
    '  4.7 行业内百分位排名',
    '  4.8 前视偏差防护',
    '5. 交易信号生成',
    '6. 策略回测系统',
    '7. 仪表盘聚合',
    '8. 报告生成引擎',
    '9. 算法参数与阈值',
    '10. 结论与局限性',
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(2)

doc.add_page_break()

# ── 1. 系统架构概述 ──
doc.add_heading('1. 系统架构概述', level=1)
doc.add_paragraph(
    '清数智算投研工作台是一套面向 A 股（沪深）和港股中长期基本面量化选股与交易信号的研究辅助系统。'
    '系统从数据采集入手，依次经过技术指标计算、五维评分、信号生成、报告撰写和策略回测，'
    '形成端到端的量化投研流水线。'
)
doc.add_paragraph('系统算法核心流程：')
steps = [
    '① 数据采集：从东方财富、Yahoo Finance、雪球等 API 获取实时行情、财务指标和券商研报；',
    '② 技术指标：基于价格序列计算 MA、MACD、RSI、布林带、成交量均线等经典技术指标；',
    '③ 五维评分：使用行业内百分位排名方法，计算质量(30)+估值(20)+成长(20)+趋势(20)+风险(10)=100分；',
    '④ 信号生成：基于评分结果和趋势确认，生成 BUY/ADD/WATCH/REDUCE/SELL 五级研究信号；',
    '⑤ 报告引擎：自动生成每日策略报告、个股深度报告和风格化策略报告；',
    '⑥ 策略回测：历史数据驱动的组合模拟，含交易成本、滑点和涨跌停约束。',
]
for s in steps:
    doc.add_paragraph(s, style='List Bullet')

doc.add_page_break()

# ── 2. 数据采集与预处理 ──
doc.add_heading('2. 数据采集与预处理', level=1)

doc.add_heading('2.1 多数据源架构', level=2)
doc.add_paragraph('系统采用 CompositeProvider 模式聚合多个数据源，按优先级自动切换：')
sources = [
    '东方财富 (EastMoneyProvider)：A 股实时行情、财务指标、券商研报、股票列表同步',
    'Yahoo Finance (YahooProvider)：港股行情、国际财务数据',
    '雪球 (XueqiuProvider)：A 股/港股补充数据',
    'AKShare：A 股宏观经济指标和市场数据',
    'MockProvider：开发/演示环境使用的模拟数据生成器',
]
for s in sources:
    doc.add_paragraph(s, style='List Bullet')

doc.add_heading('2.2 股票同步流程', level=2)
doc.add_paragraph(
    '系统启动时通过 sync_stock_list() 函数从东方财富 API 同步全量 A 股和港股列表。'
    '同步过程分页抓取（A 股约 17 页，港股约 99 页），每页 50-100 条记录，'
    '自动剔除已退市和 ST 股（保留历史数据供回测使用），入库前按 symbol 去重。'
    '数据字段包括：股票代码(symbol)、名称(name)、市场(market)、交易所(exchange)、行业(industry)、板块(sector)、状态(status)。'
)

doc.add_heading('2.3 行情数据抓取', level=2)
doc.add_paragraph(
    '使用多线程（3 个 worker）并发从数据源抓取每只股票近 180 个交易日的日线数据，'
    '包含：开盘价(open)、最高价(high)、最低价(low)、收盘价(close)、前收盘价(pre_close)、'
    '成交量(volume)、成交额(turnover)、换手率(turnover_rate)、市值(market_cap)、'
    '市盈率(PE)、市净率(PB)、股息率(dividend_yield)。'
    '每 500 只股票批量写入一次数据库以减少 IO。'
)

doc.add_heading('2.4 财务数据抓取', level=2)
doc.add_paragraph(
    '同样使用多线程（3 个 worker）抓取财务指标，包含：营收(revenue)、营收同比(revenue_yoy)、'
    '净利润(net_profit)、利润同比(net_profit_yoy)、毛利率(gross_margin)、净利率(net_margin)、'
    '净资产收益率(ROE)、总资产收益率(ROA)、资产负债率(debt_ratio)、经营现金流(operating_cashflow)、'
    '自由现金流(free_cashflow)、每股收益(EPS)、每股净资产(book_value_per_share)。'
    '财务数据按报告期(report_period)组织，支持多期追溯。'
)

doc.add_page_break()

# ── 3. 技术指标计算 ──
doc.add_heading('3. 技术指标计算', level=1)

doc.add_heading('3.1 均线系统（MA）', level=2)
doc.add_paragraph(
    '简单移动平均线 (Simple Moving Average, SMA)，计算公式：\n'
    '    MA(N) = (P₁ + P₂ + ... + P_N) / N\n'
    '系统计算 20 日(MA20)、60 日(MA60)、120 日(MA120)三条均线。'
    'MA20 反映短期趋势，MA60 反映中期趋势，MA120 反映长期趋势。'
    '数据不足对应天数时，对应均线字段为 NULL。'
)

doc.add_heading('3.2 MACD 指标', level=2)
doc.add_paragraph(
    '指数平滑异同移动平均线 (Moving Average Convergence Divergence)，计算步骤：\n'
    '    ① EMA12 = 前一日 EMA12 × 11/13 + 今日收盘价 × 2/13\n'
    '    ② EMA26 = 前一日 EMA26 × 25/27 + 今日收盘价 × 2/27\n'
    '    ③ DIF = EMA12 - EMA26\n'
    '    ④ DEA (MACD_Signal) = 前一日 DEA × 8/10 + 今日 DIF × 2/10\n'
    '    ⑤ MACD 柱 (MACD_Hist) = DIF - DEA\n'
    '初始 EMA 值使用第一日收盘价。EMA 平滑系数 α = 2/(N+1)，指数权重递减。'
)

doc.add_heading('3.3 RSI 指标', level=2)
doc.add_paragraph(
    '相对强弱指数 (Relative Strength Index, 14 日)，使用 Wilder 平滑算法：\n'
    '    ① 计算每日涨跌：gain = max(close_t - close_{t-1}, 0) / loss = max(close_{t-1} - close_t, 0)\n'
    '    ② 初始平均涨跌：AvgGain = 前 14 日 gain 均值 / AvgLoss = 前 14 日 loss 均值\n'
    '    ③ 后续使用 Wilder 平滑：AvgGain_t = (AvgGain_{t-1} × 13 + gain_t) / 14\n'
    '    ④ RS = AvgGain / AvgLoss\n'
    '    ⑤ RSI = 100 - 100 / (1 + RS)\n'
    'RSI > 70 通常视为超买，RSI < 30 通常视为超卖。数据不足 15 日时 RSI 为 NULL。'
)

doc.add_heading('3.4 布林带 (Bollinger Bands)', level=2)
doc.add_paragraph(
    '基于 20 日均线和标准差的价格通道：\n'
    '    ① 中轨 = MA20\n'
    '    ② 上轨 = MA20 + 2 × σ (σ 为 20 日收盘价总体标准差，使用 ddof=1 无偏估计)\n'
    '    ③ 下轨 = MA20 - 2 × σ\n'
    '布林带宽度 (Bandwidth) = (上轨 - 下轨) / 中轨 × 100%，用于衡量波动率。'
)

doc.add_heading('3.5 成交量均线', level=2)
doc.add_paragraph(
    '    Volume_MA5 = 近 5 日成交量均值\n'
    '    Volume_MA20 = 近 20 日成交量均值\n'
    '    成交量比率 = Volume_MA5 / Volume_MA20\n'
    '比率 1.1-2.0 视为温和放量，>2.0 视为异常放量，<1.0 视为缩量。'
)

doc.add_heading('3.6 估值补充计算', level=2)
doc.add_paragraph(
    '在种子数据阶段，从最新价格和财务数据交叉计算 PE 和 PB：\n'
    '    PE = 收盘价 / EPS（每股收益）\n'
    '    PB = 收盘价 / 每股净资产（book_value_per_share）\n'
    'EPS ≤ 0 或 book_value ≤ 0 时，对应指标为 NULL。'
)

doc.add_page_break()

# ── 4. 五维评分模型 ──
doc.add_heading('4. 五维评分模型', level=1)
doc.add_paragraph(
    '五维评分是系统的核心量化引擎，将每只股票从五个独立维度分别评分，合成 0-100 分总分。'
    '评分体系为：质量(0-30) + 估值(0-20) + 成长(0-20) + 趋势(0-20) + 风险(0-10) = 总分(0-100)。'
    '各维度之间消除指标重复计算（如 PE 只在估值维度计分，不在风险维度重复）。'
    '系统优先使用行业内百分位排名评分，当行业样本不足时回退到绝对阈值。'
)

doc.add_heading('4.1 质量评分（满分 30 分）', level=2)
doc.add_paragraph('质量维度评估公司的盈利能力和财务健康度，包含五个子指标：')
quality_items = [
    ['ROE 水平', '8 分', '行业内排名：前 20%=8 分，前 50%=6 分，前 80%=3 分，后 20%=0 分',
     '无行业数据时绝对阈值：>15%=8 分，>10%=6 分，>5%=3 分，≤5%=0 分'],
    ['ROE 趋势', '2 分', 'ROE 同比变化：提升 >2pp=2 分，稳定(-2~+2pp)=1 分，下降 >2pp=0 分', '需要前一期财务数据'],
    ['经营现金流', '8 分', '正现金流=8 分，负现金流=0 分', '极端重要的质量信号'],
    ['毛利率', '6 分', '行业内排名：前 20%=6 分，前 50%=4 分，前 80%=2 分，后 20%=0 分',
     '无行业数据时：>40%=6 分，>25%=4 分，>15%=2 分，≤15%=0 分'],
    ['资产负债率', '6 分', '行业内排名（越低越好）：前 20%（低负债）=6 分，前 50%=4 分，后 20%（高负债）=0 分',
     '无行业数据时：<40%=6 分，<60%=4 分，<75%=2 分，≥75%=0 分'],
]
table = doc.add_table(rows=len(quality_items)+1, cols=4)
table.style = 'Light Grid Accent 1'
hdr = table.rows[0].cells
hdr[0].text = '指标'; hdr[1].text = '分值'; hdr[2].text = '评分规则'; hdr[3].text = '备注'
for i, item in enumerate(quality_items):
    for j, val in enumerate(item):
        table.rows[i+1].cells[j].text = val

doc.add_heading('4.2 估值评分（满分 20 分）', level=2)
doc.add_paragraph('估值维度评估当前股价相对于基本面是否合理。本维度是 PE 的唯一计分位置，消除重复计算。')
valuation_items = [
    ['PE（市盈率）', '8 分', '行业内排名（越低越好）：前 20%=8 分，前 50%=6 分，前 80%=3 分，后 20%=0 分',
     '无行业数据：≤15=8 分，≤30=6 分，≤50=3 分，>50=0 分；PE 为负值直接 0 分'],
    ['PB（市净率）', '5 分', '行业内排名：前 20%=5 分，前 50%=3 分，后 20%=0 分',
     '无行业数据：≤2=5 分，≤5=3 分，≤10=1 分，>10=0 分'],
    ['股息率', '5 分', '行业内排名（越高越好）：前 20%（高股息）=5 分，前 50%=3 分',
     '无行业数据：>3%=5 分，>1%=3 分，≤1%=0 分'],
    ['PS（市销率）', '2 分', 'PS = 市值 / 营收；≤2=2 分，≤5=1 分，>5=0 分', '需要市值和营收数据'],
]
table = doc.add_table(rows=len(valuation_items)+1, cols=4)
table.style = 'Light Grid Accent 1'
hdr = table.rows[0].cells
hdr[0].text = '指标'; hdr[1].text = '分值'; hdr[2].text = '评分规则'; hdr[3].text = '备注'
for i, item in enumerate(valuation_items):
    for j, val in enumerate(item):
        table.rows[i+1].cells[j].text = val

doc.add_heading('4.3 成长评分（满分 20 分）', level=2)
doc.add_paragraph('成长维度评估公司收入和利润的增长动能：')
growth_items = [
    ['营收同比增长', '8 分', '>20%=8 分（高速），>10%=6 分（较快），>0%=3 分（温和），≤0%=0 分（下滑）', ''],
    ['净利润同比增长', '8 分', '>20%=8 分（高速），>10%=6 分（较快），>0%=3 分（温和），≤0%=0 分（下滑）', ''],
    ['双增长加成', '4 分', '营收和利润同时正增长=4 分；仅一项正增长=2 分；均不增长=0 分', '复合增长信号'],
]
table = doc.add_table(rows=len(growth_items)+1, cols=4)
table.style = 'Light Grid Accent 1'
hdr = table.rows[0].cells
hdr[0].text = '指标'; hdr[1].text = '分值'; hdr[2].text = '评分规则'; hdr[3].text = '备注'
for i, item in enumerate(growth_items):
    for j, val in enumerate(item):
        table.rows[i+1].cells[j].text = val

doc.add_heading('4.4 趋势评分（满分 20 分）', level=2)
doc.add_paragraph('趋势维度基于技术指标评估价格走势的健康度：')
trend_items = [
    ['收盘价 vs MA60', '6 分', '收盘价 > MA60 = 6 分（多头排列）；否则 0 分', '中期趋势判断'],
    ['MA 方向', '6 分', 'MA60 > MA120 = 6 分（上行趋势）；否则 0 分；无 MA120 时用 MA20 > MA60 替代=4 分', '趋势方向判断'],
    ['MACD 排列', '4 分', 'DIF > DEA = 4 分（多头）；MACD 柱 >0 但 DIF ≤ DEA = 2 分（偏多）；否则 0 分', ''],
    ['成交量', '4 分', '量比 1.1-2.0 = 4 分（温和放量）；>2.0 = 2 分（异常放量）；<1.0 = 0 分（缩量）', '量比=Vol_MA5/Vol_MA20'],
]
table = doc.add_table(rows=len(trend_items)+1, cols=4)
table.style = 'Light Grid Accent 1'
hdr = table.rows[0].cells
hdr[0].text = '指标'; hdr[1].text = '分值'; hdr[2].text = '评分规则'; hdr[3].text = '备注'
for i, item in enumerate(trend_items):
    for j, val in enumerate(item):
        table.rows[i+1].cells[j].text = val

doc.add_heading('4.5 风险评分（满分 10 分，分数越高风险越低）', level=2)
doc.add_paragraph('风险维度评估下行风险和不确定性：')
risk_items = [
    ['业绩稳定性', '3 分', '净利润同比 > -10% = 3 分（业绩稳定）；否则 0 分', ''],
    ['负债与现金流', '3 分', '负债率 < 70% 且经营现金流 > 0 = 3 分；负债率 < 80% = 1 分；否则 0 分', ''],
    ['波动率', '4 分', '布林带宽度 < 5% = 4 分（低波动）；5-10% = 3 分；10-15% = 1 分；>15% = 0 分', '带宽 = (上轨-下轨)/中轨×100%'],
]
table = doc.add_table(rows=len(risk_items)+1, cols=4)
table.style = 'Light Grid Accent 1'
hdr = table.rows[0].cells
hdr[0].text = '指标'; hdr[1].text = '分值'; hdr[2].text = '评分规则'; hdr[3].text = '备注'
for i, item in enumerate(risk_items):
    for j, val in enumerate(item):
        table.rows[i+1].cells[j].text = val

doc.add_heading('4.6 综合评级', level=2)
doc.add_paragraph(
    '五维总分汇总后，按以下阈值映射为五级研究评级：\n'
    '    ≥ 85 分 → BUY（高关注）：质量、估值、趋势均处于高位，建议重点研究\n'
    '    75-84 分 → ADD（增强关注）：基本面与趋势继续改善，建议增强关注\n'
    '    65-74 分 → WATCH（观察）：综合评分处于中性区间，密切关注关键指标变化\n'
    '    50-64 分 → REDUCE（风险升高）：估值或趋势出现弱化信号，风险上升\n'
    '    < 50 分 → SELL（回避观察）：基本面或风险指标显著恶化，建议回避\n'
    '研究评级是量化模型的客观输出，不构成交易指令。'
)

doc.add_heading('4.7 行业内百分位排名', level=2)
doc.add_paragraph(
    '系统在每个评分日期对所有活跃股票计算行业统计。对于每个行业，计算 PE、PB、ROE、毛利率、'
    '资产负债率、股息率、年化波动率等指标的 20/50/80 百分位值。百分位使用标准线性插值公式：\n'
    '    idx = (n - 1) × pct / 100\n'
    '    percentile = sorted_data[floor(idx)] × (1 - frac) + sorted_data[ceil(idx)] × frac\n'
    '行业内排名方法相比全市场统一阈值的优势在于：不同行业的估值中枢和盈利特征差异极大，'
    '银行天然低 PE 高杠杆，科技天然高 PE 轻资产，行业内排名能更公平地比较同行业竞争者。'
    '行业样本不足 2 只时，自动回退到绝对阈值。'
)

doc.add_heading('4.8 前视偏差防护 (Look-Ahead Bias Prevention)', level=2)
doc.add_paragraph(
    '系统在所有评分环节严格执行前视偏差防护：\n'
    '    • 价格数据：仅使用 score_date 当天或之前的最新行情（DailyPrice.trade_date ≤ score_date）\n'
    '    • 财务数据：仅使用报告期 ≤ score_date 的最新财务报告（FinancialMetric.report_period ≤ score_date）\n'
    '      考虑了财务报告的固有滞后性——Q1 报告 4 月底出，Q2 报告 8 月底出，Q3 报告 10 月底出，年报次年 4 月底出\n'
    '    • 技术指标：仅使用 trade_date ≤ score_date 的最新技术指标\n'
    '    • 评级映射：仅使用 score_date 时的可用信息决定评级\n'
    '回测中的动态评分函数 (_get_score_at_date) 进一步确保：在回测期间任意日期评分时，'
    '只使用该日期当天或之前的数据。'
)

doc.add_page_break()

# ── 5. 交易信号生成 ──
doc.add_heading('5. 交易信号生成', level=1)

doc.add_heading('5.1 信号类型决策树', level=2)
doc.add_paragraph('信号生成基于评分结果，采用多条件组合决策：')
doc.add_paragraph(
    'BUY 信号：总分 ≥ 85 且 质量 ≥ 24 且 估值 ≥ 14 且 趋势 ≥ 14\n'
    '    → 信号强度 = min(5, (总分-80)/4 + 3)，对应强度 3-5\n'
    '    → 逻辑："基本面表现较强，估值与趋势匹配，进入高关注研究名单"'
)
doc.add_paragraph(
    'ADD 信号：总分 ≥ 75 且 趋势 ≥ 12 且 风险 ≥ 7\n'
    '    → 信号强度 = min(4, (总分-70)/5 + 2)，对应强度 2-4\n'
    '    → 逻辑："基本面与趋势继续改善，建议增强关注并跟踪后续数据"'
)
doc.add_paragraph(
    'WATCH 信号：总分 ≥ 65\n'
    '    → 信号强度 = min(3, (总分-60)/5 + 1)，对应强度 1-3\n'
    '    → 逻辑：动态分析各维度短板，生成针对性观察理由'
)
doc.add_paragraph(
    'REDUCE 信号：总分 ≥ 50\n'
    '    → 信号强度 = 2\n'
    '    → 逻辑：分析估值、趋势、风险的弱化程度'
)
doc.add_paragraph(
    'SELL 信号：总分 < 50\n'
    '    → 信号强度 = 1\n'
    '    → 逻辑："基本面或风险指标明显恶化，建议回避观察"'
)

doc.add_heading('5.2 仓位计算', level=2)
doc.add_paragraph(
    '建议研究仓位基于信号类型和强度，并受行业集中度和总仓位约束：\n'
    '    基础仓位映射：BUY(5→8%, 4→6%, 3→5%) / ADD(4→5%, 3→4%, 2→3%)\n'
    '    WATCH/REDUCE/SELL → 建议仓位 0%\n'
    '    行业集中度约束：如果同行业已有仓位 > 25%，基础仓位减半（最低 1%）\n'
    '    总仓位约束：如果总仓位 > 85%，基础仓位减半（最低 1%）'
)

doc.add_heading('5.3 参考价格计算', level=2)
doc.add_paragraph(
    '入场价、目标价、止损价基于最新收盘价和布林带宽度计算：\n'
    '    波动率因子 = clamp(布林带宽度, 5%, 25%)\n'
    '    BUY/ADD：入场价 = 收盘价，目标价 = 入场价 × (1 + min(max(1.10, 1+2×波动率因子), 1.30))，止损价 = 入场价 × (1 - clamp(波动率因子, 0.88, 0.95))\n'
    '    REDUCE：止损价 = 入场价 × (1 - clamp(波动率因子×0.5, 0.90, 0.95))\n'
    '    WATCH/SELL：不计算参考价格'
)

doc.add_heading('5.4 风控信号标记', level=2)
doc.add_paragraph(
    '信号生成时自动标记风险项：风险评分 < 5、估值评分 < 10、趋势评分 < 10。'
    '这些风险标记在前端展示为橙色/红色风险标签，帮助用户快速识别潜在问题。'
)

doc.add_page_break()

# ── 6. 策略回测系统 ──
doc.add_heading('6. 策略回测系统', level=1)

doc.add_heading('6.1 回测引擎架构', level=2)
doc.add_paragraph(
    '回测系统包含两个模式：(1) 规则化策略回测 (run_backtest) 和 (2) 自定义持仓模拟 (simulate_portfolio)。'
    '两者均使用数据库中的真实历史价格运行，而非随机模拟。'
)

doc.add_heading('6.2 规则化策略回测 (run_backtest)', level=2)
doc.add_paragraph('回测流程：')
backtest_steps = [
    '① 筛选股票：选择指定市场中回测期间有行情数据的所有股票（含已退市，消除生存偏差）',
    '② 自动调整开始日期：如果最早 MA60 日期晚于用户指定开始日期，自动后延以确保技术指标可用',
    '③ 日期级步进：按交易日顺序逐日计算组合市值和基准市值',
    '④ 定期调仓：默认月度调仓（21 个交易日），可选季度调仓（63 个交易日）',
    '⑤ 评分选股：在每个调仓日，动态计算所有候选股票的评分（使用回测时点的数据，严格消除前视偏差），按评分排序，选择评分 ≥ 65 的股票，最多持仓 20 只',
    '⑥ 交易执行：卖出不在目标列表的持仓，买入目标股票，每次交易按整手（100 股）执行',
    '⑦ 涨跌停约束：涨停时无法买入，跌停时无法卖出（A 股主板 ±10%，创业板/科创板 ±20%）',
    '⑧ 风险分散：每只股票等权配置',
]
for s in backtest_steps:
    doc.add_paragraph(s, style='List Bullet')

doc.add_heading('6.3 交易成本模型', level=2)
cost_items = [
    '佣金费率：0.025%（买入和卖出均收取）',
    '印花税：0.05%（仅卖出时收取，A 股标准）',
    '过户费：0.001%（沪市收取）',
    '最低佣金：5 元/笔',
    '滑点：0.1%（买入加价 +0.1%，卖出减价 -0.1%）',
]
for c in cost_items:
    doc.add_paragraph(c, style='List Bullet')

doc.add_heading('6.4 基准构建', level=2)
doc.add_paragraph(
    '基准采用市值加权组合：以回测首日各股市值作为权重，构建买入并持有组合。'
    '无市值数据的股票赋予权重 1.0。后续每日基准市值 = Σ(初始投资 × 权重 × 当日价格 / 首日价格)。'
    '此方法近似于等权买入全市场股票，作为策略表现的比较基准。'
)

doc.add_heading('6.5 绩效指标', level=2)
metrics = [
    '总收益率 (Total Return): (最终权益 - 初始资金) / 初始资金 × 100%',
    '年化收益率 (Annual Return): (1 + 总收益率)^(1/年数) - 1',
    '基准收益率 (Benchmark Return): 市值加权基准的总收益率',
    '超额收益 (Excess Return): 策略收益 - 基准收益',
    '最大回撤 (Max Drawdown): max((峰值 - 当前值) / 峰值)，从权益曲线逐点计算',
    '夏普比率 (Sharpe Ratio): (日均收益率 × 252) / (日收益率标准差 × √252)',
    '胜率 (Win Rate): 盈利交易笔数 / 总卖出交易笔数 × 100%',
]
for m in metrics:
    doc.add_paragraph(m, style='List Bullet')

doc.add_heading('6.6 自定义持仓模拟 (simulate_portfolio)', level=2)
doc.add_paragraph(
    '允许用户自定义买入日期、股票代码和股数，系统在历史数据中查找实际买入价格，'
    '跟踪持仓至今（或指定卖出日期），计算每笔持仓的盈亏、持仓期间权益曲线和月度收益。'
    '支持同时模拟多只股票，自动处理卖出所得现金再投资。'
)

doc.add_page_break()

# ── 7. 仪表盘聚合 ──
doc.add_heading('7. 仪表盘聚合算法', level=1)

doc.add_heading('7.1 策略总结', level=2)
doc.add_paragraph(
    '仪表盘每日聚合所有交易信号，生成市场状态判断和策略总结：\n'
    '    计算 BUY+ADD 信号总数 vs REDUCE+SELL 信号总数\n'
    '    buy_add > reduce_sell × 2 → 偏多 (Bullish)，建议仓位 70-80%\n'
    '    buy_add > reduce_sell → 中性偏多 (Mildly Bullish)，建议仓位 50-70%\n'
    '    reduce_sell > buy_add × 2 → 偏空 (Bearish)，建议仓位 20-30%\n'
    '    reduce_sell > buy_add → 中性偏谨慎 (Mildly Bearish)，建议仓位 30-45%\n'
    '    其他 → 中性 (Neutral)，建议仓位 40-60%'
)

doc.add_heading('7.2 股票池分类', level=2)
doc.add_paragraph('基于评分维度的绝对分数阈值将股票自动分配到四个研究池：')
pool_items = [
    '优质基本面池：质量评分 ≥ 22 → 盈利能力突出、财务健康的股票',
    '低估值池：估值评分 ≥ 15 → 当前估值具有吸引力的股票',
    '趋势确认池：趋势评分 ≥ 15 → 技术面走势健康的股票',
    '风险预警池：风险评分 < 5 → 高风险、需要谨慎关注的股票',
]
for p in pool_items:
    doc.add_paragraph(p, style='List Bullet')

doc.add_heading('7.3 Top 信号筛选', level=2)
doc.add_paragraph(
    '仪表盘展示今日 Top 10 高关注信号（BUY/ADD），按信号强度降序排列。'
    '每条信号关联最新收盘价和日涨跌幅，帮助用户快速识别今日值得研究的标的。'
)

doc.add_heading('7.4 风险告警', level=2)
doc.add_paragraph(
    '筛选今日 REDUCE/SELL 信号，取 Top 10 作为风险告警。每条告警关联风险评分、'
    '估值评分等维度信息，标记告警级别（SELL → 高，REDUCE → 中）。'
)

doc.add_page_break()

# ── 8. 报告生成引擎 ──
doc.add_heading('8. 报告生成引擎', level=1)

doc.add_heading('8.1 每日策略报告', level=2)
doc.add_paragraph(
    'generate_daily_report() 函数从数据库获取当日所有评分和信号，聚合生成 8000+ 字专业级分析报告。'
    '报告结构包括：市场概览 → 信号统计 → 核心策略总结 → 高关注个股清单 → 行业分布热力 → 风险提示。'
    '支持三种投资风格：稳健型(steady)、进取型(aggressive)、保守型(conservative)，'
    '不同风格使用不同的评分权重和选股标准。'
)

doc.add_heading('8.2 个股深度报告', level=2)
doc.add_paragraph(
    'generate_stock_report() 为单只股票生成全面分析报告，包含：公司概览、行业地位、'
    '五维评分拆解（每项评分及理由）、技术面分析、风险因素、券商研报引用。'
    '报告内容来源于数据库中的真实评分和信号记录。'
)

doc.add_heading('8.3 风格化策略报告', level=2)
doc.add_paragraph(
    'generate_style_report() 按投资风格生成专属策略报告，三种风格的选股逻辑：\n'
    '    • 稳健型：偏好高质量(>24)、低波动(带宽<10%)、高股息(>2%）\n'
    '    • 进取型：偏好高成长(>15)、强趋势(>15)、合理估值\n'
    '    • 保守型：偏好低估值(<15 PE)、低负债(<50%)、正现金流'
)

doc.add_heading('8.4 PDF 生成', level=2)
doc.add_paragraph(
    'pdf_service.py 使用 xhtml2pdf 库将 Markdown 报告转换为 PDF，支持中文字体嵌入、'
    '分页控制、页眉页脚、水印标注（"研究辅助系统"）。PDF 下载接口受用户配额限制（默认每日 30 次）。'
)

doc.add_page_break()

# ── 9. 算法参数与阈值 ──
doc.add_heading('9. 算法参数与阈值汇总', level=1)
doc.add_paragraph('以下汇总系统中所有可调参数和阈值：')

params = [
    ['评分总分范围', '0 - 100'],
    ['质量权重 / 满分', '30% / 30 分'],
    ['估值权重 / 满分', '20% / 20 分'],
    ['成长权重 / 满分', '20% / 20 分'],
    ['趋势权重 / 满分', '20% / 20 分'],
    ['风险权重 / 满分', '10% / 10 分'],
    ['ROE 高位阈值', '> 15% (无行业数据时)'],
    ['ROE 行业前 20%', '行业内 ROE 的最高 20%'],
    ['PE 低估阈值', '≤ 15 (无行业数据时)'],
    ['PE 高估阈值', '> 50 (无行业数据时)'],
    ['PB 低估阈值', '≤ 2 (无行业数据时)'],
    ['营收高增长阈值', '> 20%'],
    ['净利润高增长阈值', '> 20%'],
    ['MACD 计算周期', 'EMA12, EMA26, Signal9'],
    ['RSI 计算周期', '14 日 (Wilder 平滑)'],
    ['布林带参数', '20 日均线, ±2 标准差'],
    ['成交量温和放大', '量比 1.1 - 2.0'],
    ['涨跌停判断', '主板 ±10%, 创业板/科创板 ±20%'],
    ['回测初始资金', '1,000,000 元（默认）'],
    ['回测调仓频率', '月度 (21 天) / 季度 (63 天)'],
    ['回测最大持仓数', '20 只'],
    ['回测选股评分下限', '≥ 65 分'],
    ['佣金费率', '0.025%'],
    ['印花税（仅卖出）', '0.05%'],
    ['滑点', '0.1%'],
    ['最低佣金', '5 元/笔'],
    ['行业集中度上限', '25% (触发仓位减半)'],
    ['总仓位上限', '85% (触发仓位减半)'],
    ['BUY 评级下限', '≥ 85 分'],
    ['ADD 评级下限', '≥ 75 分'],
    ['WATCH 评级下限', '≥ 65 分'],
    ['REDUCE 评级下限', '≥ 50 分'],
    ['仪表盘多头阈值', 'BUY+ADD > 2 × (REDUCE+SELL)'],
    ['仪表盘空头阈值', 'REDUCE+SELL > 2 × (BUY+ADD)'],
]
table = doc.add_table(rows=len(params)+1, cols=2)
table.style = 'Light Grid Accent 1'
hdr = table.rows[0].cells
hdr[0].text = '参数'; hdr[1].text = '值'
for i, (k, v) in enumerate(params):
    table.rows[i+1].cells[0].text = k
    table.rows[i+1].cells[1].text = v

doc.add_page_break()

# ── 10. 结论与局限性 ──
doc.add_heading('10. 结论与局限性', level=1)

doc.add_heading('10.1 算法优势', level=2)
strengths = [
    '多维独立评分：五维度互不重复，消除 PE 等指标的重复计分问题',
    '行业内排名：公平比较同行业竞争者，优于全市场统一阈值',
    '前视偏差防护：所有评分和回测环节严格使用历史时点数据',
    '真实交易成本：回测包含佣金、印花税、滑点、涨跌停约束',
    '生存偏差消除：回测包含已退市股票',
    '可解释性：每个评分项均有明确的业务逻辑和中文字段',
    '工程稳健性：批量查询、评分缓存、N+1 消除等性能优化',
]
for s in strengths:
    doc.add_paragraph(s, style='List Bullet')

doc.add_heading('10.2 已知局限性', level=2)
limitations = [
    '数据依赖：评分质量严重依赖财务数据的完整性和时效性。财务报告有 1-4 个月的滞后，可能不能及时反映最新经营状况',
    '行业分类：行业标签来自东方财富，可能存在分类不准确或缺失的情况（部分港股无行业标签）',
    '停牌处理：回测不处理停牌，假设停牌期间价格不变',
    '市值基准：回测基准用市值加权近似，非真正沪深 300 或恒生指数',
    '无宏观因子：当前评分模型不嵌入利率、CPI、GDP 等宏观经济变量',
    '无另类数据：未使用舆情、供应链、卫星图像等另类数据',
    '无机器学习：当前为规则型评分，未使用机器学习模型（AI 增强仅在报告生成阶段使用 DeepSeek 大模型）',
    '策略过拟合风险：评分阈值基于经验设定，在不同市场环境下可能表现不一致',
]
for l in limitations:
    doc.add_paragraph(l, style='List Bullet')

doc.add_heading('10.3 免责声明', level=2)
doc.add_paragraph(
    '本文档描述的所有算法、模型、评分、信号和回测结果仅用于学术研究和辅助分析目的，'
    '不构成任何形式的投资建议、推荐或承诺。任何基于本系统输出做出的投资决策，'
    '其风险和后果由决策者自行承担。过去的表现不代表未来的收益。'
    '使用本系统即表示您已阅读并同意上述声明。'
)

# ── 保存 ──
output_dir = os.path.dirname(os.path.abspath(__file__))
output_path = os.path.join(output_dir, '清数智算_股票分析算法技术文档.docx')
doc.save(output_path)
print(f'文档已生成: {output_path}')
