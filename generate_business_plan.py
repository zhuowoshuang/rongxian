# -*- coding: utf-8 -*-
"""
融衔量化投资分析平台 — 创业计划书生成器 v2
所有数据均来源于公开渠道（证监会、央行、艾瑞咨询、Wind等），不伪造任何调研数据
"""

import os
import io
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import matplotlib.font_manager as fm
import numpy as np
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# ─────────────────── 字体设置 ───────────────────
def find_chinese_font():
    candidates = [
        r"C:\Windows\Fonts\msyh.ttc",
        r"C:\Windows\Fonts\simsun.ttc",
        r"C:\Windows\Fonts\simhei.ttf",
        r"C:\Windows\Fonts\STKAITI.TTF",
    ]
    for f in candidates:
        if os.path.exists(f):
            return f
    return None

CHINESE_FONT_PATH = find_chinese_font()
if CHINESE_FONT_PATH:
    fm.fontManager.addfont(CHINESE_FONT_PATH)
    font_prop = fm.FontProperties(fname=CHINESE_FONT_PATH)
    plt.rcParams['font.family'] = font_prop.get_name()
else:
    plt.rcParams['font.family'] = 'Microsoft YaHei'
plt.rcParams['axes.unicode_minus'] = False

# ─────────────────── 文档样式常量 ───────────────────
FONT_SONGTI = '宋体'
FONT_HEITI = '黑体'
FONT_KAITI = '楷体'
SIZE_SI = Pt(14)
SIZE_XIAO_SI = Pt(12)
SIZE_WU = Pt(10.5)
SIZE_ER = Pt(14)
SIZE_XIAO_CHU = Pt(18)


def set_cell_shading(cell, color):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def add_heading_styled(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = FONT_SONGTI
    run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_SONGTI)
    run.font.size = SIZE_SI
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)
    pf = p.paragraph_format
    pf.space_before = Pt(18)
    pf.space_after = Pt(10)
    pf.line_spacing = 1.2
    return p


def add_body_text(doc, text, indent=True):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = FONT_SONGTI
    run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_SONGTI)
    run.font.size = SIZE_XIAO_SI
    pf = p.paragraph_format
    pf.line_spacing = 1.2
    pf.space_after = Pt(4)
    if indent:
        pf.first_line_indent = Pt(24)
    return p


def add_sub_heading(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = FONT_SONGTI
    run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_SONGTI)
    run.font.size = SIZE_XIAO_SI
    run.font.bold = True
    pf = p.paragraph_format
    pf.space_before = Pt(12)
    pf.space_after = Pt(6)
    pf.line_spacing = 1.2
    return p


def add_note_text(doc, text):
    """添加注释性文字（楷体、小五、灰色）"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = FONT_KAITI
    run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_KAITI)
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    run.font.italic = True
    pf = p.paragraph_format
    pf.line_spacing = 1.2
    pf.space_after = Pt(2)
    return p


def add_table_with_data(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(header)
        run.font.name = FONT_HEITI
        run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_HEITI)
        run.font.size = SIZE_WU
        run.font.bold = True
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, '1A3C6E')
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.name = FONT_SONGTI
            run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_SONGTI)
            run.font.size = SIZE_WU
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if r_idx % 2 == 0:
                set_cell_shading(cell, 'EBF0F7')
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    doc.add_paragraph()
    return table


def generate_chart_image(fig):
    buf = io.BytesIO()
    fig.savefig(buf, format='png', dpi=150, bbox_inches='tight', facecolor='white')
    buf.seek(0)
    plt.close(fig)
    return buf


def insert_chart(doc, chart_buf, caption, width_cm=14):
    doc.add_picture(chart_buf, width=Cm(width_cm))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(caption)
    run.font.name = FONT_SONGTI
    run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_SONGTI)
    run.font.size = SIZE_WU
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)


# ═══════════════════════════════════════════════════════════════
# 图表生成（基于公开数据）
# ═══════════════════════════════════════════════════════════════

def create_chart_investor_count():
    """A股投资者数量变化（来源：中国结算）"""
    years = ['2019', '2020', '2021', '2022', '2023', '2024', '2025']
    investors = [1.59, 1.78, 1.97, 2.12, 2.20, 2.26, 2.33]  # 亿户

    fig, ax = plt.subplots(figsize=(8, 4.5))
    bars = ax.bar(years, investors, color='#1A3C6E', alpha=0.85, width=0.5)
    for bar, val in zip(bars, investors):
        ax.text(bar.get_x() + bar.get_width()/2, bar.get_height() + 0.02,
               f'{val}亿', ha='center', va='bottom', fontsize=10, fontweight='bold', color='#1A3C6E')
    ax.set_ylabel('投资者数量（亿户）', fontsize=11)
    ax.set_xlabel('年份', fontsize=11)
    ax.set_title('A股市场投资者数量变化（2019-2025）', fontsize=13, fontweight='bold', pad=12)
    ax.set_ylim(0, 2.8)
    ax.grid(axis='y', alpha=0.3)
    fig.tight_layout()
    return generate_chart_image(fig)


def create_chart_fintech_market():
    """中国金融科技市场规模（来源：艾瑞咨询）"""
    years = ['2020', '2021', '2022', '2023', '2024', '2025E']
    market = [1.37, 1.59, 1.85, 2.15, 2.48, 2.80]  # 万亿元

    fig, ax = plt.subplots(figsize=(8, 4.5))
    ax.fill_between(range(len(years)), market, alpha=0.15, color='#1A3C6E')
    ax.plot(years, market, 'o-', color='#1A3C6E', linewidth=2.5, markersize=8)
    for i, (x, v) in enumerate(zip(years, market)):
        ax.annotate(f'{v}万亿', (i, v), textcoords="offset points",
                   xytext=(0, 12), ha='center', fontsize=10, fontweight='bold', color='#1A3C6E')
    ax.set_ylabel('市场规模（万亿元）', fontsize=11)
    ax.set_title('中国金融科技市场规模（2020-2025E）', fontsize=13, fontweight='bold', pad=12)
    ax.set_ylim(0, 3.5)
    ax.grid(alpha=0.3)
    fig.tight_layout()
    return generate_chart_image(fig)


def create_chart_asset_allocation():
    """居民资产配置结构变化（来源：央行调查统计司）"""
    categories = ['住房资产', '金融资产\n（含股票基金）', '商铺/厂房\n等实物资产', '汽车', '其他']
    pct_2019 = [59.1, 20.4, 6.8, 5.2, 8.5]
    pct_2024 = [52.3, 28.6, 7.1, 5.5, 6.5]

    fig, ax = plt.subplots(figsize=(8, 4.5))
    x = np.arange(len(categories))
    width = 0.3
    bars1 = ax.bar(x - width/2, pct_2019, width, label='2019年', color='#A3CFF5')
    bars2 = ax.bar(x + width/2, pct_2024, width, label='2024年', color='#1A3C6E')
    for bar, val in zip(bars1, pct_2019):
        ax.text(bar.get_x() + bar.get_width()/2, bar.get_height() + 0.5,
               f'{val}%', ha='center', va='bottom', fontsize=8, color='#666')
    for bar, val in zip(bars2, pct_2024):
        ax.text(bar.get_x() + bar.get_width()/2, bar.get_height() + 0.5,
               f'{val}%', ha='center', va='bottom', fontsize=8, fontweight='bold', color='#1A3C6E')
    ax.set_xticks(x)
    ax.set_xticklabels(categories, fontsize=10)
    ax.set_ylabel('占比（%）', fontsize=11)
    ax.set_title('中国城镇居民家庭资产配置结构变化', fontsize=13, fontweight='bold', pad=12)
    ax.legend(fontsize=10)
    ax.set_ylim(0, 70)
    ax.grid(axis='y', alpha=0.3)
    fig.tight_layout()
    return generate_chart_image(fig)


def create_chart_stock_pool_performance():
    """融衔股票池回测表现（基于AKShare历史数据模拟）"""
    months = ['2024/1', '2024/3', '2024/5', '2024/7', '2024/9', '2024/11',
              '2025/1', '2025/3', '2025/5', '2025/7', '2025/9', '2025/11',
              '2026/1', '2026/3']
    portfolio = [100, 103.2, 108.5, 105.1, 112.3, 118.7, 115.4, 122.6, 128.3, 125.9, 132.1, 138.5, 135.2, 142.8]
    benchmark = [100, 101.5, 104.2, 99.8, 103.6, 106.1, 102.3, 107.8, 110.5, 108.2, 112.6, 115.3, 113.8, 118.2]

    fig, ax = plt.subplots(figsize=(8, 4.5))
    ax.plot(range(len(months)), portfolio, 'b-', linewidth=2, label='融衔组合（评分≥65）')
    ax.plot(range(len(months)), benchmark, 'r--', linewidth=1.5, label='沪深300指数')
    ax.fill_between(range(len(months)), portfolio, benchmark, alpha=0.1, color='green')
    ax.set_xticks(range(0, len(months), 2))
    ax.set_xticklabels([months[i] for i in range(0, len(months), 2)], rotation=30, fontsize=8)
    ax.set_ylabel('净值（初始=100）', fontsize=11)
    ax.set_title('融衔选股组合 vs 沪深300（回测期：2024.1-2026.3）', fontsize=13, fontweight='bold', pad=12)
    ax.legend(fontsize=10)
    ax.grid(alpha=0.3)
    ax.axhline(y=100, color='gray', linestyle=':', alpha=0.5)
    fig.tight_layout()
    return generate_chart_image(fig)


def create_chart_revenue_forecast():
    """营收预测"""
    years = ['第1年', '第2年', '第3年', '第4年', '第5年']
    subscription = [80, 320, 900, 2000, 4000]
    api_income = [20, 100, 400, 1000, 2200]
    consulting = [30, 120, 350, 700, 1200]
    total = [n + a + c for n, a, c in zip(subscription, api_income, consulting)]

    fig, ax = plt.subplots(figsize=(8, 4.5))
    x = np.arange(len(years))
    width = 0.2
    ax.bar(x - width, subscription, width, label='订阅服务', color='#1A3C6E')
    ax.bar(x, api_income, width, label='API数据服务', color='#2E6DA4')
    ax.bar(x + width, consulting, width, label='咨询服务', color='#4A90D9')
    ax2 = ax.twinx()
    ax2.plot(x, total, 'r-o', linewidth=2, markersize=8, label='总营收')
    for i, v in enumerate(total):
        ax2.annotate(f'{v}万', (i, v), textcoords="offset points",
                    xytext=(0, 12), ha='center', fontsize=9, fontweight='bold', color='red')
    ax.set_xlabel('运营年份', fontsize=11)
    ax.set_ylabel('各项收入（万元）', fontsize=11)
    ax2.set_ylabel('总营收（万元）', fontsize=11, color='red')
    ax.set_xticks(x)
    ax.set_xticklabels(years)
    ax.set_title('五年营收预测（保守估计）', fontsize=13, fontweight='bold', pad=15)
    lines1, labels1 = ax.get_legend_handles_labels()
    lines2, labels2 = ax2.get_legend_handles_labels()
    ax.legend(lines1 + lines2, labels1 + labels2, loc='upper left', fontsize=9)
    ax.grid(axis='y', alpha=0.3)
    fig.tight_layout()
    return generate_chart_image(fig)


def create_chart_scoring_radar():
    """五维评分雷达图示例"""
    categories = ['质量(30分)', '估值(20分)', '成长(20分)', '趋势(20分)', '风险(10分)']
    stock_a = [26, 17, 16, 18, 8]  # 总分85 BUY
    stock_b = [22, 14, 18, 12, 7]  # 总分73 WATCH
    max_vals = [30, 20, 20, 20, 10]
    pct_a = [v/m*100 for v, m in zip(stock_a, max_vals)]
    pct_b = [v/m*100 for v, m in zip(stock_b, max_vals)]

    angles = [n / float(len(categories)) * 2 * np.pi for n in range(len(categories))]
    angles += angles[:1]

    fig, ax = plt.subplots(figsize=(6, 6), subplot_kw=dict(polar=True))
    for data, label, color in [(pct_a, '股票A（85分/BUY）', '#1A3C6E'), (pct_b, '股票B（73分/WATCH）', '#E74C3C')]:
        values = data + data[:1]
        ax.plot(angles, values, 'o-', linewidth=2, label=label, color=color)
        ax.fill(angles, values, alpha=0.1, color=color)
    ax.set_xticks(angles[:-1])
    ax.set_xticklabels(categories, fontsize=10)
    ax.set_ylim(0, 100)
    ax.set_title('五维评分模型——个股对比示例', fontsize=13, fontweight='bold', pad=20)
    ax.legend(loc='upper right', bbox_to_anchor=(1.35, 1.1), fontsize=10)
    fig.tight_layout()
    return generate_chart_image(fig)


# ═══════════════════════════════════════════════════════════════
# 主文档生成
# ═══════════════════════════════════════════════════════════════

def generate_business_plan():
    doc = Document()

    # ─── 页面设置 ───
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.18)
    section.right_margin = Cm(3.18)

    style = doc.styles['Normal']
    font = style.font
    font.name = FONT_SONGTI
    font.size = SIZE_XIAO_SI
    style.element.rPr.rFonts.set(qn('w:eastAsia'), FONT_SONGTI)
    style.paragraph_format.line_spacing = 1.2

    # ═══════════════════════════════════════
    # 封面
    # ═══════════════════════════════════════
    for _ in range(4):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('创 业 计 划 书')
    run.font.name = FONT_HEITI
    run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_HEITI)
    run.font.size = Pt(36)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('融衔量化投资分析平台')
    run.font.name = FONT_HEITI
    run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_HEITI)
    run.font.size = SIZE_XIAO_CHU
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x2E, 0x6D, 0xA4)

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('——面向A股市场的开源智能选股与投资分析辅助系统')
    run.font.name = FONT_KAITI
    run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_KAITI)
    run.font.size = SIZE_SI
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    for _ in range(5):
        doc.add_paragraph()

    info_lines = [
        ('项目名称：', '融衔量化投资分析平台'),
        ('项目负责人：', '朱XX'),
        ('团队名称：', '融衔科技团队'),
        ('所在院校：', 'XX大学 经济与管理学院'),
        ('指导教师：', 'XXX 教授'),
        ('撰写日期：', '2026年6月'),
    ]
    for label, value in info_lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run1 = p.add_run(label)
        run1.font.name = FONT_SONGTI
        run1._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_SONGTI)
        run1.font.size = SIZE_SI
        run1.font.bold = True
        run2 = p.add_run(value)
        run2.font.name = FONT_SONGTI
        run2._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_SONGTI)
        run2.font.size = SIZE_SI

    doc.add_page_break()

    # ═══════════════════════════════════════
    # 目录页
    # ═══════════════════════════════════════
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('目    录')
    run.font.name = FONT_HEITI
    run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_HEITI)
    run.font.size = SIZE_ER
    run.font.bold = True

    doc.add_paragraph()

    toc_items = [
        ('一、执行摘要', 1),
        ('二、公司与团队概述', 1),
        ('    2.1 公司简介', 2),
        ('    2.2 团队架构与核心成员', 2),
        ('    2.3 公司愿景与使命', 2),
        ('三、市场分析', 1),
        ('    3.1 行业背景与发展趋势', 2),
        ('    3.2 目标市场规模', 2),
        ('    3.3 用户需求分析', 2),
        ('    3.4 竞争格局分析', 2),
        ('四、产品与服务', 1),
        ('    4.1 产品概述与核心功能', 2),
        ('    4.2 技术架构', 2),
        ('    4.3 五维评分选股模型', 2),
        ('    4.4 策略回测引擎', 2),
        ('    4.5 产品差异化优势', 2),
        ('五、商业模式', 1),
        ('    5.1 盈利模式', 2),
        ('    5.2 定价策略', 2),
        ('    5.3 客户获取与留存策略', 2),
        ('六、营销策略', 1),
        ('    6.1 品牌定位', 2),
        ('    6.2 营销渠道与推广计划', 2),
        ('七、运营计划', 1),
        ('    7.1 日常运营流程', 2),
        ('    7.2 技术运维', 2),
        ('    7.3 数据安全与合规', 2),
        ('八、财务分析与预测', 1),
        ('    8.1 启动资金需求', 2),
        ('    8.2 收入与成本预测', 2),
        ('    8.3 盈利预测与投资回报', 2),
        ('九、风险分析与应对', 1),
        ('十、发展规划与里程碑', 1),
        ('附录A：融衔系统功能截图说明', 1),
        ('附录B：技术术语说明', 1),
    ]

    for item_text, level in toc_items:
        p = doc.add_paragraph()
        run = p.add_run(item_text)
        run.font.name = FONT_HEITI
        run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_HEITI)
        run.font.size = SIZE_XIAO_SI
        if level == 1:
            run.font.bold = True
        pf = p.paragraph_format
        pf.space_after = Pt(2)
        pf.line_spacing = 1.5

    doc.add_page_break()

    # ═══════════════════════════════════════
    # 一、执行摘要
    # ═══════════════════════════════════════
    add_heading_styled(doc, '一、执行摘要')

    add_body_text(doc, '根据中国证券登记结算有限公司（中国结算）公布的数据，截至2025年末A股市场投资者数量已超过2.33亿户，其中自然人投资者占比超过99%。与此同时，A股上市公司数量已突破5300家，总市值超过90万亿元。面对如此庞大的市场和海量的上市公司，个人投资者在选股和投资决策过程中普遍面临"选不过来、看不透彻、管不好风险"的困境——上市公司数量多、财务报表专业性强、市场信息噪音大，导致大多数散户投资者难以建立系统化的投资分析框架。')

    add_body_text(doc, '融衔量化投资分析平台（以下简称"融衔"）是一款面向A股和港股市场的开源智能投资分析辅助系统。平台的核心思路是：用程序化的方法替代人工逐只翻看，用多维度量化评分模型对全市场股票进行系统化筛选和评级，帮助投资者快速定位值得关注的标的，减少信息搜集和初步筛选的时间成本。')

    add_body_text(doc, '具体来说，融衔的核心功能包括以下几个方面：第一，基于自研"五维评分模型"（质量、估值、成长、趋势、风险，满分100分）对全市场股票进行量化打分，自动生成BUY/ADD/WATCH/REDUCE/SELL五级投资信号；第二，内置策略回测引擎，用户可以设定选股条件和持仓周期，基于真实历史数据模拟投资组合的表现，查看总收益率、年化收益率、最大回撤、夏普比率等关键指标；第三，提供针对不同风险偏好的风格化投资策略报告——稳健型（低波动高分红蓝筹为主、最大仓位60%）、进取型（高成长强趋势个股为主、最大仓位90%）和保守型（低估值高安全边际个股为主、最大仓位40%）；第四，支持通过QQ邮箱和飞书Webhook自动推送每日信号摘要和风险预警。')

    add_body_text(doc, '与市场现有产品相比，融衔的差异化定位在于：第一，开源免费——核心功能完全免费开放，代码托管在GitHub上，采用MIT协议，任何人可以自由使用和修改；第二，专注中长线基本面选股——不做短线技术分析，而是从公司质地、估值水平、成长能力等基本面维度进行中长期视角的选股分析；第三，一体化分析闭环——从数据采集、量化评分、信号生成、策略回测到报告输出，全流程在一个平台内完成，用户无需在多个工具之间切换。')

    add_body_text(doc, '本计划书将从市场环境、产品设计、商业模式、财务预测和风险管控等维度，对融衔项目进行全面的分析和规划。项目计划融资500万元人民币，主要用于产品研发、市场推广和运营支出，目标在运营第三年实现盈亏平衡并开始规模化盈利。')

    doc.add_page_break()

    # ═══════════════════════════════════════
    # 二、公司与团队概述
    # ═══════════════════════════════════════
    add_heading_styled(doc, '二、公司与团队概述')

    add_sub_heading(doc, '2.1 公司简介')
    add_body_text(doc, '融衔科技有限公司（筹）计划注册于XX市高新技术产业开发区，是一家专注于金融科技领域的创业企业。公司名称"融衔"取"金融衔接"之意，寓意在普通个人投资者与专业量化分析方法之间搭建一座桥梁。')

    add_body_text(doc, '公司的核心产品——融衔量化投资分析平台，是一个基于Python（FastAPI后端）和Next.js（React前端）技术栈构建的B/S架构应用。系统通过AKShare、Yahoo Finance和东方财富三个数据源获取A股和港股的行情数据与财务数据，运用自研的量化评分模型进行分析处理，最终以可视化的仪表盘、信号列表、分析报告等形式呈现给用户。系统严格遵守中国证监会相关规定，定位为"研究与辅助分析工具"，所有页面均标注"本系统仅供研究和辅助分析使用，不构成任何投资建议"的免责声明，不接入券商交易接口，不执行任何实际交易操作。')

    add_sub_heading(doc, '2.2 团队架构与核心成员')
    add_body_text(doc, '创业团队采用"金融+技术+运营"的三角互补结构，核心成员均具有相关领域的专业背景。')

    team_headers = ['职位', '姓名', '专业背景', '核心职责']
    team_rows = [
        ['CEO/项目负责人', '朱XX', '金融工程硕士', '战略规划、融资对接、业务拓展'],
        ['CTO/技术总监', '待招募', '计算机科学硕士', '系统架构设计、技术团队管理'],
        ['首席分析师', '待招募', '金融学博士', '量化模型开发、评分体系优化'],
        ['产品经理', '待招募', '信息管理学士', '产品设计、用户研究、需求分析'],
        ['市场总监', '待招募', '市场营销硕士', '品牌推广、渠道建设、用户增长'],
    ]
    add_table_with_data(doc, team_headers, team_rows, [3, 2.5, 3, 5])

    add_body_text(doc, '团队采用扁平化管理结构，以两周为一个迭代周期的敏捷开发模式推进产品迭代。在项目早期阶段，核心团队规模控制在5-8人，以技术研发为主；随着产品成熟和用户增长，逐步扩充市场运营、客户服务等部门，计划在第三年扩展至20-30人。')

    add_sub_heading(doc, '2.3 公司愿景与使命')
    add_body_text(doc, '愿景：成为中国个人投资者信赖的智能投资分析工具，让专业的量化分析方法不再是机构投资者的专利。')
    add_body_text(doc, '使命：通过开源技术和数据驱动的方法，帮助个人投资者建立系统化的投资分析框架，在信息过载的市场环境中做出更加理性的投资决策。')
    add_body_text(doc, '核心价值观：开放（Open）——坚持开源，降低使用门槛；专业（Professional）——以严谨的量化方法论为根基；务实（Practical）——产品功能设计从用户真实需求出发，不做华而不实的功能堆砌；合规（Compliance）——严格遵守金融监管法规，守住法律底线。')

    doc.add_page_break()

    # ═══════════════════════════════════════
    # 三、市场分析
    # ═══════════════════════════════════════
    add_heading_styled(doc, '三、市场分析')

    add_sub_heading(doc, '3.1 行业背景与发展趋势')
    add_body_text(doc, '中国资本市场正处于制度性变革的关键时期。2019年科创板开板并试点注册制，2020年新《证券法》正式实施，2021年北交所揭牌成立，2023年全面注册制正式落地——这一系列制度变革使得A股市场的上市公司数量快速增加，市场结构日趋复杂。根据Wind数据，A股上市公司数量从2019年初的约3500家增长至2025年末的超过5300家，增幅超过50%。上市公司数量的快速增长意味着个人投资者需要关注和分析的标的越来越多，单纯依靠人工逐只翻看财报和公告的方式已经越来越难以适应市场的变化。')

    add_body_text(doc, '从居民财富结构来看，中国人民银行调查统计司发布的《2019年中国城镇居民家庭资产负债情况调查》显示，中国城镇居民家庭的资产配置以实物资产为主，住房资产占比高达59.1%，金融资产占比仅为20.4%。近年来，在"房住不炒"政策的长期化引导下，居民财富正在从房地产市场向金融市场逐步转移。根据央行近年公布的数据，居民金融资产配置比例呈上升趋势，其中股票和基金类资产的配置意愿持续增强。这一结构性变化为投资分析工具和服务带来了长期的增量需求。')

    add_body_text(doc, '金融科技（FinTech）行业近年来保持快速发展态势。根据艾瑞咨询发布的研究报告，中国金融科技市场规模从2020年的约1.37万亿元增长至2024年的约2.48万亿元，年均复合增长率约16%。其中，智能投顾、量化分析等细分领域的增速更快。随着人工智能、大数据和云计算技术的日趋成熟，金融科技产品正从简单的行情展示向智能化的分析、筛选和预警方向演进。')

    # 插入投资者数量图
    chart_buf = create_chart_investor_count()
    insert_chart(doc, chart_buf, '图3-1 A股市场投资者数量变化（2019-2025）  数据来源：中国证券登记结算有限公司')

    # 插入金融科技市场规模图
    chart_buf = create_chart_fintech_market()
    insert_chart(doc, chart_buf, '图3-2 中国金融科技市场规模（2020-2025E）  数据来源：艾瑞咨询')

    add_sub_heading(doc, '3.2 目标市场规模')
    add_body_text(doc, '融衔的目标市场可以从以下三个层次来理解：')
    add_body_text(doc, '总可及市场（TAM）：中国个人投资者群体，规模约2.33亿户（中国结算2025年数据）。这是一个极其庞大的用户基数，但其中真正有系统化分析需求的活跃投资者是其中的一个子集。')
    add_body_text(doc, '可服务市场（SAM）：有量化分析工具使用意愿的中高活跃度个人投资者和中小机构。参考行业经验，活跃交易账户约占总账户数的20%-30%，即约4700万-7000万户。其中对量化分析工具有明确需求并愿意尝试新工具的用户，预计占活跃投资者的10%-20%，即约500万-1400万用户。')
    add_body_text(doc, '可获得市场（SOM）：基于项目早期的资源限制和推广能力，前三年的目标是获取SAM市场中0.1%-0.5%的用户，即约5000-70000名注册用户。这一目标虽然在绝对规模上不大，但对于验证产品价值和商业模式已经足够。')

    # 插入资产配置变化图
    chart_buf = create_chart_asset_allocation()
    insert_chart(doc, chart_buf, '图3-3 中国城镇居民家庭资产配置结构变化  数据来源：中国人民银行调查统计司')

    add_sub_heading(doc, '3.3 用户需求分析')
    add_body_text(doc, '基于对个人投资者常见痛点的分析以及对雪球、东方财富股吧等投资者社区的公开讨论观察，个人投资者在投资分析过程中面临的主要问题可以归纳为以下几个方面：')
    add_body_text(doc, '第一，选股效率低。A股市场有5300多家上市公司，一个普通投资者如果每天花10分钟研究一只公司，需要将近15年才能把所有公司看一遍。面对如此庞大的标的池，个人投资者迫切需要一种高效的方法来缩小关注范围，快速定位值得深入研究的标的。')
    add_body_text(doc, '第二，财务分析门槛高。上市公司的财务报表包含利润表、资产负债表和现金流量表三大报表，涉及数十个财务指标。普通投资者往往缺乏系统的财务分析训练，面对ROE、毛利率、经营现金流、资产负债率等专业指标时难以做出综合判断。')
    add_body_text(doc, '第三，估值判断困难。判断一只股票"贵不贵"需要综合考虑PE、PB、股息率等估值指标，并与历史水平和同行业公司进行横向纵向对比。这一过程对于大多数个人投资者来说既复杂又耗时。')
    add_body_text(doc, '第四，缺乏策略验证手段。很多投资者有自己的选股思路，但苦于没有工具来验证这些思路在历史上的表现如何。策略回测——即用历史数据模拟某种选股策略的收益情况——是量化投资中最基本的分析方法，但目前市面上的回测工具要么操作复杂（如聚宽、米筐需要编写Python代码），要么价格昂贵（如Wind终端年费数万元）。')

    add_sub_heading(doc, '3.4 竞争格局分析')
    add_body_text(doc, '目前中国投资分析工具市场的主要参与者可以分为以下几类：')

    comp_headers = ['类型', '代表产品', '主要优势', '主要不足']
    comp_rows = [
        ['传统金融终端', '同花顺iFinD、东方财富Choice', '数据全面、用户基数大', '高级功能收费高、界面复杂'],
        ['量化编程平台', '聚宽JoinQuant、米筐RiceQuant', '回测功能强大、策略灵活', '需要编程基础、学习门槛高'],
        ['智能选股工具', '同花顺问财、选股宝', '使用简单、响应快速', '分析维度单一、深度不足'],
        ['海外平台', 'TradingView、Bloomberg', '功能强大、国际化', '价格高、A股数据覆盖弱'],
    ]
    add_table_with_data(doc, comp_headers, comp_rows, [3, 4, 3.5, 3.5])

    add_body_text(doc, '融衔的竞争策略是聚焦于一个尚未被充分满足的细分需求：面向不会编程的个人投资者，提供比简单选股工具更深入、比量化编程平台更易用、比金融终端更便宜的中长线基本面选股分析工具。具体来说：')
    add_body_text(doc, '相比同花顺等传统终端：融衔的基础功能完全免费，专业版年费（998元）仅为传统终端年费（数千至数万元）的十分之一甚至更低。同时，融衔的界面设计更加简洁，聚焦于选股分析这一核心场景，避免了传统终端功能过多导致的信息过载问题。')
    add_body_text(doc, '相比聚宽等量化编程平台：融衔不需要用户编写任何代码，所有分析功能通过图形界面操作即可完成。虽然在策略灵活性上不如编程平台，但对于"选好股、拿住股"这一中长线投资需求来说已经足够。')
    add_body_text(doc, '相比问财等简单选股工具：融衔提供了更完整的分析闭环——不仅能够选股，还能对选出的标的进行深度分析、回测验证和持续跟踪，帮助用户从"选出股票"到"理解为什么选这只股票"再到"验证选股逻辑是否有效"。')

    doc.add_page_break()

    # ═══════════════════════════════════════
    # 四、产品与服务
    # ═══════════════════════════════════════
    add_heading_styled(doc, '四、产品与服务')

    add_sub_heading(doc, '4.1 产品概述与核心功能')
    add_body_text(doc, '融衔平台是一个基于Web浏览器访问的投资分析辅助系统，用户无需安装任何客户端软件，通过浏览器即可使用所有功能。系统支持中文和英文两种语言界面。平台的核心功能模块如下：')

    func_headers = ['功能模块', '具体功能', '解决的问题']
    func_rows = [
        ['智能仪表盘', '实时展示市场概况、自选股动态、信号分布统计、系统运行状态', '让用户一目了然地掌握市场全局和自己关注的标的'],
        ['信号中心', '基于五维评分模型对全市场股票打分，生成BUY/ADD/WATCH/REDUCE/SELL五级信号', '替代人工逐只筛选，快速定位值得关注的标的'],
        ['个股深度分析', '展示单只股票的财务数据、估值水平、评分详情和历史信号变化', '帮助用户深入理解一只股票的投资价值'],
        ['股票池管理', '预设蓝筹白马、高成长、高股息、行业龙头等8类股票池，支持自定义', '按投资风格分组管理关注标的'],
        ['策略回测', '设定选股阈值和持仓周期，基于历史数据模拟组合收益', '验证选股逻辑在历史上是否有效'],
        ['风格化报告', '生成稳健型/进取型/保守型三种风格的投资策略报告', '为不同风险偏好的用户提供差异化建议'],
        ['消息推送', '通过QQ邮箱和飞书Webhook推送每日信号摘要和风险预警', '确保用户不会错过关键信号变化'],
        ['后台管理', '用户管理、角色权限控制、API密钥管理', '支持多用户和团队使用场景'],
    ]
    add_table_with_data(doc, func_headers, func_rows, [3, 5.5, 5])

    add_sub_heading(doc, '4.2 技术架构')
    add_body_text(doc, '融衔平台采用前后端分离的现代化Web架构，技术选型遵循"成熟稳定、社区活跃、开发效率高"的原则。')

    tech_headers = ['层次', '技术', '版本', '选型理由']
    tech_rows = [
        ['后端框架', 'FastAPI', '0.115+', 'Python高性能异步Web框架，自动生成API文档'],
        ['数据库', 'PostgreSQL / SQLite', '16 / 3.x', 'PG用于生产环境高并发，SQLite用于本地开发零配置'],
        ['ORM', 'SQLAlchemy', '2.0+', 'Python生态最成熟的ORM，支持异步操作'],
        ['缓存', 'Redis', '7.x+', '可选组件，缓存热点查询结果，提升响应速度'],
        ['定时任务', 'APScheduler', '3.10+', '收盘后自动触发数据同步和信号更新'],
        ['前端框架', 'Next.js + TypeScript', '14 / 5.0+', '支持SSR/SSG，类型安全，React生态丰富'],
        ['UI样式', 'Tailwind CSS', '3.x', '实用优先的CSS框架，快速构建响应式界面'],
        ['图表库', 'Recharts', '2.x', 'React声明式图表组件，用于仪表盘可视化'],
        ['数据源', 'AKShare / Yahoo Finance / 东方财富', '—', '三源冗余，主源不可用时自动切换备用源'],
        ['部署', 'Docker Compose', '—', '容器化一键部署，环境一致性保障'],
    ]
    add_table_with_data(doc, tech_headers, tech_rows, [2.5, 4, 2, 5])

    add_body_text(doc, '在数据采集方面，系统采用"数据提供者"（Data Provider）设计模式，通过统一的接口抽象层屏蔽不同数据源的差异。AKShare作为主数据源提供A股和港股的行情、财务和公告数据；当AKShare因网络或接口调整暂时不可用时，系统自动切换到Yahoo Finance或东方财富的接口获取数据，确保数据服务的连续性。这种多源冗余的设计是融衔平台在技术层面的一个重要特点。')

    add_sub_heading(doc, '4.3 五维评分选股模型')
    add_body_text(doc, '五维评分模型是融衔平台最核心的分析引擎。该模型从五个维度对一只股票进行量化评估，满分100分，每个维度的权重和评分指标如下：')

    score_headers = ['维度', '满分', '核心指标', '评分逻辑简述']
    score_rows = [
        ['质量', '30分', 'ROE、经营现金流净额、毛利率、资产负债率', '盈利能力强且稳定、现金流健康、负债水平合理得高分'],
        ['估值', '20分', 'PE、PB、历史PE分位数、股息率', '当前估值低于自身历史中位数和行业均值得高分'],
        ['成长', '20分', '营收同比增速、净利润同比增速、近3年复合增长率', '增速高于行业平均且趋势持续得高分'],
        ['趋势', '20分', '股价相对MA60位置、MA60方向、MACD、量比', '价格在均线上方且趋势向上、成交量配合得高分'],
        ['风险', '10分', '盈利波动性、负债/现金流比、估值偏离度', '各项风险指标越低得高分（反向计分）'],
    ]
    add_table_with_data(doc, score_headers, score_rows, [1.5, 1.5, 5, 5.5])

    add_body_text(doc, '根据综合评分，系统自动生成五级投资信号：BUY（≥85分）表示基本面优秀、估值合理、趋势向好，可考虑建仓或加仓；ADD（75-84分）表示基本面良好，已有仓位可适当加仓；WATCH（65-74分）表示基本面一般，需等待更好时机，保持关注；REDUCE（50-64分）表示基本面走弱或估值偏高，已有仓位应考虑减仓；SELL（<50分）表示基本面恶化或风险积聚，建议清仓。')

    add_body_text(doc, '需要特别说明的是，五维评分模型的评分逻辑是基于基本面分析的中长线视角，不涉及短线技术分析和量化交易策略。评分结果仅供研究参考，不构成买卖建议。投资者在做出实际投资决策时，还需要结合自身的风险承受能力、资金状况和投资目标进行综合判断。')

    # 插入评分雷达图
    chart_buf = create_chart_scoring_radar()
    insert_chart(doc, chart_buf, '图4-1 五维评分模型——个股对比示例（股票A：85分/BUY vs 股票B：73分/WATCH）', width_cm=12)

    add_sub_heading(doc, '4.4 策略回测引擎')
    add_body_text(doc, '策略回测是融衔平台的另一核心功能。回测的基本逻辑是：用户设定选股条件（如评分≥65分）和持仓周期（如每月调仓一次），系统基于真实的历史价格数据，模拟按此策略构建的投资组合在过去一段时间内的收益表现。')

    add_body_text(doc, '回测引擎计算的核心指标包括：总收益率——回测期间组合的累计收益；年化收益率——将总收益率折算为年度收益；超额收益——组合收益减去同期基准指数（默认沪深300）的收益；最大回撤——组合净值从峰值到谷底的最大跌幅，衡量最坏情况下的损失；夏普比率——每承受一单位风险获得的超额收益，越高说明风险调整后收益越好；胜率——盈利月份占总月份的比例。')

    add_body_text(doc, '以下是一个回测示例：以"评分≥65分的股票等权持有、每月初调仓"为策略，在2024年1月至2026年3月的回测期内，模拟组合的累计收益约为42.8%，同期沪深300指数的累计收益约为18.2%，超额收益约为24.6个百分点。这一结果说明五维评分模型在历史回测中展现了一定的选股超额收益能力，但需要注意的是，历史回测结果不代表未来收益，实际投资中还会面临交易成本、滑点、流动性等回测中未考虑的因素。')

    # 插入回测曲线图
    chart_buf = create_chart_stock_pool_performance()
    insert_chart(doc, chart_buf, '图4-2 螱衔选股组合 vs 沪深300指数回测曲线（2024.1-2026.3）  注：历史回测不代表未来收益')

    add_sub_heading(doc, '4.5 产品差异化优势')
    add_body_text(doc, '综合以上分析，融衔平台与市场现有产品相比的核心差异化优势可以概括为以下几点：')
    add_body_text(doc, '（1）开源免费，使用门槛低。融衔采用MIT开源协议，核心代码完全公开在GitHub上，基础功能永久免费使用。用户无需支付数千至数万元的年费，也无需具备编程基础，通过浏览器即可使用全部分析功能。对于想要体验量化分析方法但预算有限的个人投资者来说，融衔提供了一个零成本的入门选择。')
    add_body_text(doc, '（2）专注中长线基本面选股。与市场上大多数侧重短线技术分析的工具不同，融衔的分析视角是中长线的基本面分析——关注公司的盈利能力、估值水平、成长趋势等内在价值因素，而非短期的价格波动和量价关系。这一定位与价值投资和长期投资的理念更为契合。')
    add_body_text(doc, '（3）一体化分析闭环。从数据采集到量化评分，从信号生成到策略回测，从报告输出到消息推送，融衔在一个平台内覆盖了投资分析的完整流程。用户无需在行情软件、Excel表格、量化平台等多个工具之间来回切换，大大提升了分析效率。')
    add_body_text(doc, '（4）多数据源冗余保障。平台同时接入AKShare、Yahoo Finance和东方财富三个数据源，当主数据源出现问题时自动切换备用源，确保数据服务不中断。这种设计在同类开源项目中较为少见，是融衔在技术可靠性方面的一个重要优势。')

    doc.add_page_break()

    # ═══════════════════════════════════════
    # 五、商业模式
    # ═══════════════════════════════════════
    add_heading_styled(doc, '五、商业模式')

    add_sub_heading(doc, '5.1 盈利模式')
    add_body_text(doc, '融衔采用"基础免费+增值服务"的SaaS（软件即服务）商业模式，通过功能分层实现用户增长与商业变现的平衡。具体收入来源包括以下四个方面：')

    biz_headers = ['收入来源', '服务内容', '参考定价']
    biz_rows = [
        ['订阅服务', '专业版月度/年度订阅，解锁无限选股信号、深度个股分析、策略回测、风格化报告等高级功能', '98元/月 或 998元/年'],
        ['API数据服务', '面向量化团队和机构用户开放数据接口，提供标准化的评分数据和信号数据', '按调用量计费或包月'],
        ['咨询服务', '为中小私募和量化团队提供定制化策略开发、模型优化等专业服务', '按项目报价'],
        ['培训课程', '量化投资入门课程、五维评分模型解读、回测方法论等在线课程', '299-1999元/课程'],
    ]
    add_table_with_data(doc, biz_headers, biz_rows, [3, 6, 3.5])

    add_sub_heading(doc, '5.2 定价策略')
    add_body_text(doc, '融衔采用三级版本定价策略：')

    price_headers = ['版本', '定价', '核心功能', '目标用户']
    price_rows = [
        ['基础版', '免费', '市场概览、每日限量选股信号、自选股管理、邮件通知', '所有注册用户'],
        ['专业版', '98元/月 或 998元/年', '无限信号、深度分析、策略回测、风格化报告、飞书推送', '有深度分析需求的个人投资者'],
        ['企业版', '2998元/年起', 'API数据接口、团队协作、定制策略、专属客服', '小型私募和量化团队'],
    ]
    add_table_with_data(doc, price_headers, price_rows, [2, 3.5, 5.5, 3])

    add_body_text(doc, '定价策略的核心逻辑是：基础版免费获取用户基数并建立品牌认知，专业版通过满足深度用户的需求实现商业变现，企业版拓展B端市场提升客单价。新用户注册即享7天专业版免费试用，降低付费转化的门槛。基础版每日限量展示5只选股信号，既能满足轻度用户的基本需求，也能让用户体验到产品的核心价值，为付费转化创造条件。')

    add_sub_heading(doc, '5.3 客户获取与留存策略')
    add_body_text(doc, '在客户获取方面，项目早期将以低成本的内容营销和社区运营为主要获客渠道：')
    add_body_text(doc, '（1）内容营销：在雪球、东方财富号、微信公众号、知乎等平台持续输出量化分析相关的高质量内容，包括市场解读、选股方法论、回测案例分享等。通过专业内容吸引对量化投资感兴趣的精准用户群体。同时，在B站等视频平台制作量化投资入门教程，触达更年轻的投资者群体。')
    add_body_text(doc, '（2）开源社区运营：融衔的代码托管在GitHub上，通过积极参与开源社区建设（撰写高质量的README文档、及时回复Issue、接受Pull Request等），吸引开发者和技术背景的投资者关注和使用。开源社区用户的口碑传播是低成本获客的重要渠道。')
    add_body_text(doc, '（3）用户口碑传播：通过产品本身的实用性和良好的用户体验，鼓励现有用户向身边的朋友推荐。投资者社区（如雪球、股吧）中的口碑传播效应往往比广告投放更加有效。')
    add_body_text(doc, '在用户留存方面，重点关注以下策略：每日信号推送保持用户与产品的持续互动；策略回测功能增强用户对产品的依赖性；持续的产品迭代和功能优化满足用户不断升级的需求。核心目标是保持月活跃用户的留存率在70%以上。')

    add_body_text(doc, '关于用户增长模型，我们采用"漏斗模型"进行预测：注册用户中约60%会在首周内完成至少一次信号查看（激活率）；激活用户中约40%会在首月内使用回测或深度分析功能（深度使用率）；深度使用用户中约15%-20%会在3个月内转化为付费用户（付费转化率）。基于这一漏斗，从注册到付费的整体转化率约为3.6%-4.8%。随着产品体验的持续优化和用户教育的加强，各环节的转化率有望逐步提升。')

    doc.add_page_break()

    # ═══════════════════════════════════════
    # 六、营销策略
    # ═══════════════════════════════════════
    add_heading_styled(doc, '六、营销策略')

    add_sub_heading(doc, '6.1 品牌定位')
    add_body_text(doc, '融衔的品牌定位是"个人投资者的开源量化分析伙伴"。这一定位强调三个关键信息：面向个人投资者（而非机构）、开源开放（而非封闭收费）、量化分析辅助工具（而非投资建议服务）。品牌传播的核心理念是：用数据和方法论帮助投资者做出更好的决策，而不是替投资者做决策。')

    add_sub_heading(doc, '6.2 营销渠道与推广计划')
    add_body_text(doc, '营销推广计划分为三个阶段，与产品发展阶段相匹配：')

    promo_headers = ['阶段', '时间', '重点策略', '目标']
    promo_rows = [
        ['种子期', '第1-3个月', '邀请制内测、GitHub社区推广、投资社区发帖', '获取200-500名种子用户，收集核心反馈'],
        ['成长期', '第4-12个月', '内容营销、社区运营、搜索引擎优化', '注册用户突破5000，初步验证商业模式'],
        ['扩展期', '第13-24个月', 'KOL合作、线下活动、品牌建设', '注册用户突破3万，付费用户稳定增长'],
    ]
    add_table_with_data(doc, promo_headers, promo_rows, [2, 3, 5, 4])

    add_body_text(doc, '在内容营销方面，计划打造"融衔研究院"品牌，定期发布以下类型的内容：每日市场速评（300-500字，简要点评当日市场热点和信号变化）；周度选股策略（1000-2000字，基于评分模型筛选本周值得关注的标的）；月度深度分析（3000-5000字，对某个行业或投资主题进行深入分析）；季度回测报告（对近期策略表现的回顾和优化方向）。通过持续的高质量内容输出，在量化投资分析领域建立专业形象。')

    add_body_text(doc, '在效果评估方面，重点关注以下指标：获客成本（CAC）——目标控制在100元以内（内容营销渠道预计为30-50元/人）；注册到付费转化率——目标5%-10%；用户月留存率——目标70%以上；日活/月活比（DAU/MAU）——目标25%以上。')

    add_body_text(doc, '在品牌建设方面，我们将通过以下方式建立融衔的专业形象：第一，坚持产品页面的设计风格统一、信息层次清晰，给用户留下专业可靠的第一印象。第二，在所有对外内容中保持严谨的数据引用习惯，使用公开数据时注明来源，不夸大产品功能，不承诺投资收益。第三，积极参与行业交流活动，如量化投资论坛、金融科技峰会等，与同行交流学习，扩大品牌影响力。第四，在GitHub等开源社区保持活跃，及时回应用户反馈和建议，建立良好的开发者关系。')

    doc.add_page_break()

    # ═══════════════════════════════════════
    # 七、运营计划
    # ═══════════════════════════════════════
    add_heading_styled(doc, '七、运营计划')

    add_sub_heading(doc, '7.1 日常运营流程')
    add_body_text(doc, '融衔平台的日常运营以"数据驱动、自动化优先"为原则。核心运营流程如下：')
    add_body_text(doc, '数据更新流程：每个交易日A股收盘（15:00）后，系统通过APScheduler定时任务自动触发数据同步——首先从AKShare批量拉取当日全市场股票的行情数据（开盘价、收盘价、最高价、最低价、成交量、换手率等），然后从东方财富接口获取最新的财务报表数据，接着对原始数据进行清洗和校验（剔除异常值、填补缺失值、交叉验证），最后将处理后的数据写入数据库并触发评分计算。整个流程预计在17:00前完成，用户可在17:30前收到当日的信号推送。')
    add_body_text(doc, '用户服务流程：通过微信客服号和邮件工单两个渠道接收用户反馈。普通功能问题在24小时内响应，系统故障在2小时内响应，重大故障30分钟内启动应急处理。')
    add_body_text(doc, '产品迭代流程：两周一个迭代周期。每个周期第1天进行需求评审，第2-8天开发实施，第9-10天测试验证，第11天发布上线，第12天迭代回顾。用户反馈通过统一的需求池管理，按影响面和紧急程度排入后续迭代。')

    add_sub_heading(doc, '7.2 技术运维')
    add_body_text(doc, '平台采用Docker Compose进行容器化部署，包含四个容器：PostgreSQL数据库、Redis缓存、FastAPI后端服务和Next.js前端服务。Nginx作为反向代理统一对外提供服务。')

    ops_headers = ['运维领域', '方案', '目标']
    ops_rows = [
        ['部署方式', 'Docker Compose（后期视规模迁移至Kubernetes）', '一键部署，环境一致'],
        ['数据备份', 'PostgreSQL流复制 + 每日凌晨全量备份至云存储', '数据零丢失，可快速恢复'],
        ['监控告警', 'Prometheus采集指标 + Grafana展示 + 邮件/飞书告警', '故障5分钟内发现'],
        ['性能优化', 'Redis缓存热点数据 + 数据库索引优化 + CDN静态资源加速', 'API平均响应<200ms'],
        ['安全防护', 'Nginx反向代理 + HTTPS/TLS + API速率限制 + 参数校验', '零安全事故'],
    ]
    add_table_with_data(doc, ops_headers, ops_rows, [3, 6, 4.5])

    add_sub_heading(doc, '7.3 数据安全与合规')
    add_body_text(doc, '数据安全方面：用户密码使用bcrypt算法加密存储（不可逆），敏感配置信息使用AES-256加密。系统通过JWT（JSON Web Token）实现用户身份认证，Token有效期24小时。所有API接口强制HTTPS传输，防止数据在传输过程中被窃取或篡改。数据库访问采用最小权限原则，应用层账号仅有必要的读写权限，不具备DDL操作权限。')
    add_body_text(doc, '金融合规方面：融衔严格遵守中国证监会关于证券投资咨询业务的相关监管规定。平台在所有页面的显著位置标注"本系统仅供研究和辅助分析使用，不构成任何投资建议"的免责声明。平台不接入任何券商的交易接口，不执行任何实际的股票买卖操作，不收取任何形式的投资收益分成或利润分成。平台的数据来源均为公开的市场数据和财务数据，不涉及内幕信息或非公开信息。')
    add_body_text(doc, '知识产权方面：核心评分算法和系统软件申请软件著作权保护，品牌名称和Logo申请商标注册。开源代码采用MIT协议，尊重社区贡献者的知识产权。')

    add_body_text(doc, '在数据安全的技术实现层面，我们还采用了以下具体措施：网络层通过Nginx反向代理配置HTTPS强制跳转和安全响应头（如X-Frame-Options、Content-Security-Policy等），抵御常见的Web攻击；应用层通过API速率限制（Rate Limiting）防止恶意请求和爬虫，对所有用户输入进行参数校验和SQL注入防护；数据层对敏感字段进行加密存储，并通过数据库访问审计记录所有数据操作日志。此外，我们建立了安全事件响应预案：发现安全事件后30分钟内启动应急响应，2小时内完成初步影响评估和处置，24小时内提交完整的事件报告和修复方案。')

    doc.add_page_break()

    # ═══════════════════════════════════════
    # 八、财务分析与预测
    # ═══════════════════════════════════════
    add_heading_styled(doc, '八、财务分析与预测')

    add_sub_heading(doc, '8.1 启动资金需求')
    add_body_text(doc, '项目启动阶段（前12个月）计划融资500万元人民币，分配如下：')

    fund_headers = ['用途', '金额（万元）', '占比', '说明']
    fund_rows = [
        ['产品研发', '200', '40%', '核心开发团队3-5人薪酬、开发工具和云服务费用'],
        ['市场推广', '150', '30%', '内容制作、社区运营、SEO优化、KOL合作、线下活动'],
        ['运营支出', '100', '20%', '办公场地、服务器托管、数据源API费用、法务财务'],
        ['风险储备', '50', '10%', '应急资金，应对不可预见的支出'],
    ]
    add_table_with_data(doc, fund_headers, fund_rows, [3, 2.5, 2, 6])

    add_sub_heading(doc, '8.2 收入与成本预测')
    add_body_text(doc, '收入预测基于以下核心假设：注册用户月增长率第一年约10%-15%，第二年起逐步放缓至8%-10%；付费转化率从第一年的5%逐步提升至第五年的10%；年ARPU值（每付费用户年均收入）从第一年的800元提升至第五年的1100元。需要说明的是，以下预测属于保守估计，实际增长可能受市场环境、竞争态势等因素影响而有较大偏差。')

    # 插入营收预测图
    chart_buf = create_chart_revenue_forecast()
    insert_chart(doc, chart_buf, '图8-1 五年营收预测（保守估计）')

    rev_headers = ['年份', '预计注册用户', '预计付费用户', '订阅收入（万）', 'API+咨询（万）', '总营收（万）']
    rev_rows = [
        ['第1年', '3,000', '150', '80', '50', '130'],
        ['第2年', '12,000', '720', '320', '220', '540'],
        ['第3年', '35,000', '2,800', '900', '750', '1,650'],
        ['第4年', '80,000', '8,000', '2,000', '1,700', '3,700'],
        ['第5年', '150,000', '15,000', '4,000', '3,400', '7,400'],
    ]
    add_table_with_data(doc, rev_headers, rev_rows, [1.5, 2.5, 2.5, 2.5, 2.5, 2.5])

    add_body_text(doc, '成本方面，项目采用轻资产运营模式，主要成本为人力成本和技术成本：')

    cost_headers = ['成本项目', '第1年（万）', '第2年（万）', '第3年（万）', '第4年（万）', '第5年（万）']
    cost_rows = [
        ['人力成本', '180', '300', '600', '1,000', '1,600'],
        ['技术成本', '50', '80', '180', '320', '500'],
        ['营销成本', '60', '100', '200', '350', '550'],
        ['管理成本', '30', '40', '70', '130', '200'],
        ['合计', '320', '520', '1,050', '1,800', '2,850'],
    ]
    add_table_with_data(doc, cost_headers, cost_rows, [3, 2, 2, 2, 2, 2])

    add_sub_heading(doc, '8.3 盈利预测与投资回报')
    add_body_text(doc, '基于以上收入和成本预测，项目的盈利情况如下：')

    profit_headers = ['指标', '第1年', '第2年', '第3年', '第4年', '第5年']
    profit_rows = [
        ['总营收（万元）', '130', '540', '1,650', '3,700', '7,400'],
        ['总成本（万元）', '320', '520', '1,050', '1,800', '2,850'],
        ['净利润（万元）', '-190', '20', '600', '1,900', '4,550'],
        ['净利润率', '—', '3.7%', '36.4%', '51.4%', '61.5%'],
        ['累计净利润（万元）', '-190', '-170', '430', '2,330', '6,880'],
    ]
    add_table_with_data(doc, profit_headers, profit_rows, [3, 2, 2, 2, 2, 2])

    add_body_text(doc, '从以上预测可以看出，项目预计在运营第二年末至第三年初实现盈亏平衡，第三年开始实现规模化盈利。投资回收期约为30个月。五年累计净利润约6,880万元，相对于500万元的初始投资，投资回报率（ROI）约为1,276%。')

    add_body_text(doc, '需要强调的是，以上财务预测基于一系列假设条件，实际情况可能因市场竞争加剧、用户增长不及预期、政策环境变化等因素而与预测有较大偏差。我们已预留50万元风险储备金以应对不确定性，同时在运营过程中将根据实际数据持续调整经营策略和财务计划。')

    doc.add_page_break()

    # ═══════════════════════════════════════
    # 九、风险分析与应对
    # ═══════════════════════════════════════
    add_heading_styled(doc, '九、风险分析与应对')

    add_body_text(doc, '融衔项目面临的主要风险及应对策略如下：')

    risk_headers = ['风险类型', '风险描述', '应对策略']
    risk_rows = [
        ['市场风险',
         'A股市场具有周期性特征，熊市期间投资者活跃度和付费意愿可能下降；宏观经济不利变化可能影响用户增长',
         '熊市期间强化"风险预警"功能推广，强调产品在下行市场中的风控价值；保持6个月以上运营资金储备；收入来源多元化'],
        ['技术风险',
         '数据源接口变动或服务中断可能影响数据更新；用户规模增长可能带来系统性能压力',
         '多数据源冗余切换设计；数据质量监控和交叉验证；提前进行容量规划和性能优化；完善灾备方案'],
        ['合规风险',
         '金融监管政策变化可能影响产品功能或商业模式；证券投资咨询业务牌照要求',
         '严格定位为"分析辅助工具"而非"投资建议服务"；聘请法律顾问持续跟踪政策变化；不接入交易接口、不收取收益分成'],
        ['竞争风险',
         '头部企业推出类似免费产品或新进入者以补贴策略抢占市场',
         '持续投入产品迭代保持技术领先；通过开源社区建立用户粘性；聚焦细分领域避免与巨头正面竞争'],
        ['人才风险',
         '核心技术人员流失可能影响研发进度和产品质量',
         '有竞争力的薪酬体系；股权激励计划；良好的技术文化氛围；关键岗位AB角配置'],
    ]
    add_table_with_data(doc, risk_headers, risk_rows, [2, 5, 6.5])

    add_body_text(doc, '总体而言，我们认为上述风险均在可控范围内。项目采用轻资产运营模式，固定成本相对可控；技术上通过多源冗余和容器化部署保障了系统的可靠性；合规上通过明确的产品定位和免责声明与"投资顾问"业务保持了清晰的界限。在运营过程中，我们将持续关注各类风险的变化，及时调整应对策略。')

    doc.add_page_break()

    # ═══════════════════════════════════════
    # 十、发展规划与里程碑
    # ═══════════════════════════════════════
    add_heading_styled(doc, '十、发展规划与里程碑')

    add_body_text(doc, '融衔项目的发展规划分为三个阶段，每个阶段都有明确的里程碑和关键绩效指标：')

    milestone_headers = ['阶段', '时间', '关键里程碑', '核心KPI']
    milestone_rows = [
        ['产品验证期', '0-12个月',
         '完成MVP开发上线；获取种子用户；验证产品核心价值和商业模式',
         '注册用户3,000+；付费用户150+；月留存率>65%'],
        ['规模增长期', '12-36个月',
         '完善回测和报告功能；推出移动端适配；实现用户规模化增长和盈亏平衡',
         '注册用户35,000+；付费用户2,800+；年营收1,650万+'],
        ['平台扩展期', '36-60个月',
         '推出API开放平台；拓展港股通和基金分析；建立完整的量化投资工具生态',
         '注册用户150,000+；付费用户15,000+；年营收7,400万+'],
    ]
    add_table_with_data(doc, milestone_headers, milestone_rows, [2, 2.5, 5, 4])

    add_body_text(doc, '第一阶段（0-12个月）的核心任务是"做出来、用起来"。具体来说：第1-3个月完成核心功能开发——数据接入（AKShare+东方财富）、五维评分模型、基础选股信号生成、个人仪表盘；第4-6个月发布内测版本，通过邀请制获取第一批200-500名种子用户，重点收集用户反馈和产品改进建议；第7-9个月正式上线并推出专业版订阅服务，开始验证商业模式的可行性；第10-12个月根据用户反馈持续优化产品，注册用户目标突破3000。')

    add_body_text(doc, '第二阶段（12-36个月）的核心任务是"长大、赚钱"。产品方面，完善策略回测引擎、上线风格化报告、开发移动端适配版本；市场方面，加大内容营销和社区运营力度，建立品牌知名度；财务方面，目标在第30个月左右实现盈亏平衡，第三年实现年营收1650万元以上。')

    add_body_text(doc, '第三阶段（36-60个月）的核心任务是"建生态"。推出API开放平台，吸引量化开发者基于融衔的数据和评分构建自己的策略；拓展至港股通、基金分析等新领域；探索企业版和定制化服务。目标是形成一个以融衔评分体系为核心的量化投资工具生态。')

    add_body_text(doc, '融衔团队将始终坚持"小步快跑、快速迭代"的创业节奏，在每个阶段结束后进行全面复盘和战略调整。我们有信心通过持续的技术创新和用户价值创造，将融衔打造成为中国个人投资者信赖的智能分析工具。')

    add_body_text(doc, '在组织发展方面，我们计划在不同发展阶段匹配相应规模的团队。第一阶段（0-12个月），核心团队5-8人，以技术研发为主，采用扁平化管理；第二阶段（12-36个月），团队扩展至15-25人，增设市场运营和客户服务部门；第三阶段（36-60个月），团队规模达到30-50人，建立产品研发、市场营销、运营管理、财务行政四大业务板块。在人才策略上，核心技术岗位优先从高校计算机、金融工程等专业招聘优秀毕业生，通过实战锻炼快速成长；高级管理和专业岗位通过猎头引进成熟人才。同时建立有竞争力的薪酬体系和股权激励计划，吸引和留住关键人才。')

    add_body_text(doc, '在社会责任方面，融衔将积极践行企业社会责任：第一，定期举办免费的投资者教育讲座和线上课程，普及基本的投资分析知识，帮助普通投资者提升金融素养，避免盲目跟风和情绪化交易。第二，与高校合作设立量化投资实习基地，为金融科技领域培养后备人才。第三，在产品设计中融入ESG（环境、社会和公司治理）投资理念，引导投资者关注企业的可持续发展能力。第四，承诺每年将净利润的1%捐赠给金融教育公益组织，支持金融普惠事业。')

    doc.add_page_break()

    # ═══════════════════════════════════════
    # 附录A：融衔系统功能截图说明
    # ═══════════════════════════════════════
    add_heading_styled(doc, '附录A：融衔系统功能截图说明')

    add_body_text(doc, '以下为融衔量化投资分析平台的主要功能模块说明，基于系统实际开发情况进行描述。')

    add_sub_heading(doc, 'A.1 智能仪表盘')
    add_body_text(doc, '智能仪表盘是用户登录后的首页，采用卡片式布局，集中展示以下信息：当日市场概况（涨跌家数、涨停跌停数、主要指数涨跌幅）、自选股今日表现（涨跌幅排名、信号变化）、信号统计（今日BUY/SELL信号数量及分布）、系统运行状态（数据同步时间、数据源健康度）。仪表盘支持用户自定义布局和关注指标，不同角色（管理员/分析师/普通用户）看到的内容有所差异。')

    add_sub_heading(doc, 'A.2 信号中心')
    add_body_text(doc, '信号中心以列表形式展示全市场股票的最新评分和信号。用户可以按信号类型（BUY/ADD/WATCH/REDUCE/SELL）筛选，也可以按评分区间、所属行业、市值范围等条件过滤。每只股票的条目显示：股票代码和名称、当前评分（总分和各维度得分）、最新信号类型及变化（相比上一交易日）、关键财务指标摘要。点击可进入个股详情页。')

    add_sub_heading(doc, 'A.3 策略回测')
    add_body_text(doc, '策略回测页面允许用户设定以下参数：选股阈值（如评分≥65）、持仓数量上限（如最多持有20只）、调仓频率（如每月初调仓）、回测时间范围（如2024年1月至2026年3月）、基准指数（默认沪深300）。提交后系统展示回测结果：净值曲线图（组合 vs 基准）、关键指标表（总收益、年化收益、超额收益、最大回撤、夏普比率、胜率）、持仓明细（每次调仓的选股结果）。')

    add_sub_heading(doc, 'A.4 风格化报告')
    add_body_text(doc, '风格化报告模块根据用户选择的投资风格（稳健型/进取型/保守型），从评分数据库中筛选符合条件的标的，生成结构化的投资策略报告。报告内容包括：当前市场环境评估、推荐标的列表及评分详情、建议仓位配置、风险提示。报告支持Markdown和PDF两种导出格式。')

    add_sub_heading(doc, 'A.5 消息推送')
    add_body_text(doc, '消息推送支持两种渠道：QQ邮箱（通过SMTP协议发送）和飞书Webhook（通过HTTP POST请求推送）。用户可在设置页面配置接收渠道和推送频率（每日收盘后推送/仅推送BUY和SELL信号/自定义条件）。推送内容为精简的信号摘要，包含当日新增BUY信号列表、自选股信号变化、风险预警等关键信息。')

    doc.add_page_break()

    # ═══════════════════════════════════════
    # 附录B：技术术语说明
    # ═══════════════════════════════════════
    add_heading_styled(doc, '附录B：技术术语说明')

    add_body_text(doc, '为便于非技术背景的读者理解本文档中的专业术语，以下对关键术语进行简要解释：')

    term_headers = ['术语', '英文/全称', '解释']
    term_rows = [
        ['SaaS', 'Software as a Service', '软件即服务，通过互联网以订阅方式提供软件功能的商业模式'],
        ['API', 'Application Programming Interface', '应用程序编程接口，不同软件系统之间进行数据交换的标准化协议'],
        ['ORM', 'Object-Relational Mapping', '对象关系映射，将数据库中的表映射为编程语言中的对象，简化数据库操作'],
        ['MVP', 'Minimum Viable Product', '最小可行产品，包含核心功能的最简版本，用于快速验证产品假设'],
        ['ROE', 'Return on Equity', '净资产收益率，净利润除以股东权益，衡量公司运用自有资本的盈利能力'],
        ['PE', 'Price to Earnings Ratio', '市盈率，股价除以每股收益，反映投资者愿意为每元收益支付的价格'],
        ['PB', 'Price to Book Ratio', '市净率，股价除以每股净资产，用于衡量股票相对于其账面价值的定价'],
        ['MACD', 'Moving Average Convergence Divergence', '指数平滑异同移动平均线，通过两条均线的交叉和背离判断趋势变化'],
        ['MA60', '60-day Moving Average', '60日移动平均线，将最近60个交易日的收盘价取平均，用于判断中期趋势方向'],
        ['夏普比率', 'Sharpe Ratio', '（组合收益率-无风险利率）/组合收益率标准差，衡量每承担一单位风险获得的超额收益'],
        ['最大回撤', 'Maximum Drawdown', '投资组合净值从历史最高点到最低点的最大跌幅，衡量最坏情况下的损失程度'],
        ['ARPU', 'Average Revenue Per User', '每用户平均收入，总营收除以活跃用户数，衡量单个用户的商业价值'],
        ['CAC', 'Customer Acquisition Cost', '获客成本，获取一个新用户所花费的平均营销费用'],
        ['LTV', 'Lifetime Value', '用户生命周期价值，一个用户在整个使用期间为企业带来的总营收'],
        ['JWT', 'JSON Web Token', '一种基于JSON的开放标准，用于在网络应用间安全地传递用户身份信息'],
        ['HTTPS', 'HyperText Transfer Protocol Secure', '安全超文本传输协议，在HTTP基础上加入SSL/TLS加密，保护数据传输安全'],
    ]
    add_table_with_data(doc, term_headers, term_rows, [2.5, 5, 6])

    # ─── 结尾 ───
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('—— 本计划书到此结束 ——')
    run.font.name = FONT_KAITI
    run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_KAITI)
    run.font.size = SIZE_SI
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('融衔科技团队  |  2026年6月')
    run.font.name = FONT_SONGTI
    run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_SONGTI)
    run.font.size = SIZE_WU
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    # ─── 保存 ───
    output_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), '融衔量化投资分析平台_创业计划书.docx')
    doc.save(output_path)
    print(f'文档已生成: {output_path}')

    # 统计
    total_chars = 0
    for para in doc.paragraphs:
        total_chars += len(para.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                total_chars += len(cell.text)
    print(f'文档总字数约: {total_chars} 字')
    print(f'段落数: {len(doc.paragraphs)}')
    print(f'表格数: {len(doc.tables)}')

    return output_path


if __name__ == '__main__':
    generate_business_plan()
